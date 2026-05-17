"""
generate_summary_tables.py
~~~~~~~~~~~~~~~~~~~~~~~~~~
Standalone module for generating formatted summary tables from
autocorrelation and wavelet feature comparison results.

Produces:
  - CSV:  full feature table (all features, median / mean±SD / IQR / p / delta)
  - DOCX: formatted Word table with highlighted significant rows

Usage (called from compare_autocorrelation.py and visualize_wavelets.py):

    from generate_summary_tables import save_autocorrelation_tables, save_wavelet_tables

    save_autocorrelation_tables(
        comparison_rows=comparison_rows,
        labeled_rows=labeled_rows,
        output_dir=OUTPUT_DIR,
    )

    save_wavelet_tables(
        labeled_rows=labeled_rows,
        comparison_rows=comparison_rows,
        output_dir=OUTPUT_DIR,
    )

Requirements: python-docx  (pip install python-docx)
"""

from __future__ import annotations

import csv
from pathlib import Path

import numpy as np
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def _iqr(values: np.ndarray) -> tuple[float, float]:
    """Return (Q1, Q3)."""
    return float(np.percentile(values, 25)), float(np.percentile(values, 75))


def _get_values(rows: list[dict], feature: str, group: str) -> np.ndarray:
    return np.array(
        [float(r[feature]) for r in rows if r["group"] == group],
        dtype=np.float64,
    )


def _save_csv(rows: list[dict], path: Path) -> None:
    if not rows:
        print(f"[SKIP] No rows to save: {path}")
        return
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(
    cell,
    text: str,
    bold: bool = False,
    size: int = 9,
    center: bool = False,
    color: RGBColor | None = None,
) -> None:
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Autocorrelation tables ────────────────────────────────────────────────────

def save_autocorrelation_tables(
    comparison_rows: list[dict],
    labeled_rows: list[dict],
    output_dir: Path,
) -> None:
    """
    Save autocorrelation feature summary as CSV and formatted DOCX.

    CSV columns: Feature | Control median | Control mean±SD | Patient median |
                 Patient mean±SD | n_control | n_patient | p-value | Cliff's δ | Significant
    DOCX: median [IQR] per group, significant rows highlighted in yellow.
    """
    _save_autocorrelation_csv(comparison_rows, output_dir / "autocorrelation_summary_table.csv")
    _save_autocorrelation_docx(
        comparison_rows=comparison_rows,
        labeled_rows=labeled_rows,
        out_path=output_dir / "autocorrelation_summary_table.docx",
    )


def _save_autocorrelation_csv(comparison_rows: list[dict], out_path: Path) -> None:
    table_rows = []
    for r in comparison_rows:
        table_rows.append({
            "Feature":              r["feature"],
            "Control median":       f"{r['control_median']:.4f}",
            "Control mean \u00b1 SD": f"{r['control_mean']:.4f} \u00b1 {r['control_std']:.4f}",
            "Patient median":       f"{r['patient_median']:.4f}",
            "Patient mean \u00b1 SD":  f"{r['patient_mean']:.4f} \u00b1 {r['patient_std']:.4f}",
            "n control":            r["n_control"],
            "n patient":            r["n_patient"],
            "p-value":              f"{r['p_value']:.4f}",
            "Cliff's delta":        f"{r['cliffs_delta_patient_vs_control']:.3f}",
            "Significant (p<0.05)": "Yes" if r["p_value"] < 0.05 else "No",
        })
    _save_csv(table_rows, out_path)
    print(f"[OK] Autocorrelation summary CSV saved: {out_path.name}")


def _save_autocorrelation_docx(
    comparison_rows: list[dict],
    labeled_rows: list[dict],
    out_path: Path,
) -> None:
    # pre-compute IQR — feature names come from comparison_rows itself
    iqr_data: dict[str, dict] = {}
    for r in comparison_rows:
        feat = r["feature"]
        pat  = _get_values(labeled_rows, feat, "patient")
        ctrl = _get_values(labeled_rows, feat, "control")
        p_q1, p_q3 = _iqr(pat)  if len(pat)  > 0 else (0.0, 0.0)
        c_q1, c_q3 = _iqr(ctrl) if len(ctrl) > 0 else (0.0, 0.0)
        iqr_data[feat] = {"pat_q1": p_q1, "pat_q3": p_q3,
                          "ctrl_q1": c_q1, "ctrl_q3": c_q3}

    doc = DocxDocument()

    tp = doc.add_paragraph()
    tr = tp.add_run("Autocorrelation Feature Comparison: Patients vs Controls")
    tr.bold = True
    tr.font.size = Pt(13)

    cap = doc.add_paragraph(
        "Values reported as median [Q1\u2013Q3]. "
        "p-values from Wilcoxon rank-sum test (two-sided). "
        "Cliff\u2019s \u03b4: negative = patients lower than controls. "
        "Highlighted rows: p\u00a0<\u00a00.05."
    )
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    headers = ["Feature", "Control median [IQR]", "Patient median [IQR]",
               "p-value", "Cliff\u2019s \u03b4", "Significant"]
    col_widths = [Inches(1.6), Inches(1.9), Inches(1.9),
                  Inches(0.85), Inches(0.85), Inches(0.9)]

    table = doc.add_table(rows=1 + len(comparison_rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        _cell_text(hdr.cells[i], h, bold=True, size=9, center=True,
                   color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_bg(hdr.cells[i], "2E75B6")

    for ri, r in enumerate(comparison_rows):
        feat = r["feature"]
        iq   = iqr_data.get(feat, {})
        sig  = r["p_value"] < 0.05
        bg   = "FFF2CC" if sig else "FFFFFF"
        dr   = table.rows[ri + 1]
        vals = [
            feat,
            f"{r['control_median']:.4f} [{iq.get('ctrl_q1', 0):.4f}\u2013{iq.get('ctrl_q3', 0):.4f}]",
            f"{r['patient_median']:.4f} [{iq.get('pat_q1', 0):.4f}\u2013{iq.get('pat_q3', 0):.4f}]",
            f"{r['p_value']:.4f}",
            f"{r['cliffs_delta_patient_vs_control']:.3f}",
            "Yes" if sig else "No",
        ]
        for i, val in enumerate(vals):
            _cell_text(dr.cells[i], val, bold=(i == 0), size=9, center=(i > 0))
            _set_cell_bg(dr.cells[i], bg)

    doc.save(str(out_path))
    print(f"[OK] Autocorrelation summary DOCX saved: {out_path.name}")


# ── Wavelet tables ────────────────────────────────────────────────────────────

SUBBAND_MAP = {
    "w01": "HH lv1", "w02": "LH lv1", "w03": "HL lv1",
    "w04": "HH lv2", "w05": "LH lv2", "w06": "HL lv2",
    "w07": "HH lv3", "w08": "LH lv3", "w09": "HL lv3",
    "w10": "Residual",
}
QUADRANT_MAP = {
    "tl": "Top Left", "tr": "Top Right",
    "bl": "Bottom Left", "br": "Bottom Right",
}
QUAD_COLORS = {
    "tl": "EBF5FB", "tr": "EAF4F4",
    "bl": "FEF9E7", "br": "FDEDEC",
}


def save_wavelet_tables(
    labeled_rows: list[dict],
    comparison_rows: list[dict],
    output_dir: Path,
) -> None:
    """
    Save wavelet feature summaries:
      - Full CSV: all features with median [IQR], p-value, Cliff's delta
      - Significant DOCX: only p < 0.05 features, colour-coded by quadrant
    """
    _save_wavelet_csv(labeled_rows, comparison_rows, output_dir / "wavelet_summary_table_full.csv")
    _save_wavelet_docx(labeled_rows, comparison_rows, output_dir / "wavelet_summary_table_significant.docx")


def _save_wavelet_csv(
    labeled_rows: list[dict],
    comparison_rows: list[dict],
    out_path: Path,
) -> None:
    table_rows = []
    for r in comparison_rows:
        feat = r["feature"]
        pat  = _get_values(labeled_rows, feat, "patient")
        ctrl = _get_values(labeled_rows, feat, "control")
        p_q1, p_q3 = _iqr(pat)  if len(pat)  > 0 else (0.0, 0.0)
        c_q1, c_q3 = _iqr(ctrl) if len(ctrl) > 0 else (0.0, 0.0)
        parts    = feat.split("_")
        quadrant = QUADRANT_MAP.get(parts[0], parts[0]) if len(parts) > 1 else ""
        subband  = SUBBAND_MAP.get(parts[1], parts[1])  if len(parts) > 1 else feat
        table_rows.append({
            "Feature":              feat,
            "Quadrant":             quadrant,
            "Subband":              subband,
            "Control median":       f"{r['control_median']:.6f}",
            "Control IQR":          f"[{c_q1:.6f}\u2013{c_q3:.6f}]",
            "Patient median":       f"{r['patient_median']:.6f}",
            "Patient IQR":          f"[{p_q1:.6f}\u2013{p_q3:.6f}]",
            "n_control":            r["n_control"],
            "n_patient":            r["n_patient"],
            "p-value":              f"{r['p_value']:.4f}",
            "Cliff's delta":        f"{r['cliffs_delta_patient_vs_control']:.3f}",
            "Significant (p<0.05)": "Yes" if r["p_value"] < 0.05 else "No",
        })
    _save_csv(table_rows, out_path)
    print(f"[OK] Full wavelet summary CSV saved: {out_path.name}")


def _save_wavelet_docx(
    labeled_rows: list[dict],
    comparison_rows: list[dict],
    out_path: Path,
) -> None:
    sig_rows = [r for r in comparison_rows if r["p_value"] < 0.05]
    if not sig_rows:
        print("[SKIP] No significant wavelet features — skipping DOCX table.")
        return

    doc = DocxDocument()

    tp = doc.add_paragraph()
    tr = tp.add_run(
        f"Wavelet Feature Comparison: Patients vs Controls "
        f"({len(sig_rows)} significant features, p\u00a0<\u00a00.05)"
    )
    tr.bold = True
    tr.font.size = Pt(13)

    cap = doc.add_paragraph(
        "Values reported as median [Q1\u2013Q3]. "
        "p-values from Wilcoxon rank-sum test (two-sided). "
        "Cliff\u2019s \u03b4: negative = patients lower than controls. "
        "Row colour indicates quadrant: "
        "blue\u2009=\u2009Top Left, green\u2009=\u2009Top Right, "
        "yellow\u2009=\u2009Bottom Left, red\u2009=\u2009Bottom Right."
    )
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    headers = ["Feature", "Quadrant", "Subband",
               "Control median [IQR]", "Patient median [IQR]",
               "p-value", "Cliff\u2019s \u03b4"]
    col_widths = [Inches(1.0), Inches(0.95), Inches(0.85),
                  Inches(1.75), Inches(1.75), Inches(0.8), Inches(0.8)]

    table = doc.add_table(rows=1 + len(sig_rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        _cell_text(hdr.cells[i], h, bold=True, size=9, center=True,
                   color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_bg(hdr.cells[i], "1F3864")

    for ri, r in enumerate(sig_rows):
        feat   = r["feature"]
        parts  = feat.split("_")
        q_key  = parts[0] if len(parts) > 1 else ""
        w_key  = parts[1] if len(parts) > 1 else feat
        pat    = _get_values(labeled_rows, feat, "patient")
        ctrl   = _get_values(labeled_rows, feat, "control")
        p_q1, p_q3 = _iqr(pat)  if len(pat)  > 0 else (0.0, 0.0)
        c_q1, c_q3 = _iqr(ctrl) if len(ctrl) > 0 else (0.0, 0.0)
        bg     = QUAD_COLORS.get(q_key, "FFFFFF")
        dr     = table.rows[ri + 1]
        vals   = [
            feat,
            QUADRANT_MAP.get(q_key, q_key),
            SUBBAND_MAP.get(w_key, w_key),
            f"{r['control_median']:.4f} [{c_q1:.4f}\u2013{c_q3:.4f}]",
            f"{r['patient_median']:.4f} [{p_q1:.4f}\u2013{p_q3:.4f}]",
            f"{r['p_value']:.4f}",
            f"{r['cliffs_delta_patient_vs_control']:.3f}",
        ]
        for i, val in enumerate(vals):
            _cell_text(dr.cells[i], val, bold=(i == 0), size=9, center=(i > 0))
            _set_cell_bg(dr.cells[i], bg)

    doc.save(str(out_path))
    print(f"[OK] Significant wavelet features DOCX saved: {out_path.name}")


# ── Standalone entry point ────────────────────────────────────────────────────

if __name__ == "__main__":
    print(
        "generate_summary_tables.py is a library module.\n"
        "Import save_autocorrelation_tables or save_wavelet_tables from it.\n"
        "Run compare_autocorrelation.py or visualize_wavelets.py to generate tables."
    )